from pathlib import Path

from docx import Document

_TEMPLATE = Path(__file__).parent / "templates" / "report_template.docx"


def _add_content(doc, text: str):
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith(("- ", "• ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif len(line) > 2 and line[0].isdigit() and line[1] in ".)" and line[2] == " ":
            doc.add_paragraph(line[3:], style="List Number")
        else:
            doc.add_paragraph(line)


def create_report(
    title: str,
    content: str,
    filename: str,
    sections: str | None = None,
    output_path: str | None = None,
) -> str:
    
    base = Path(output_path) if output_path else Path("output")
    base.mkdir(parents=True, exist_ok=True)
    output = base / f"{filename}.docx"

    doc = Document(_TEMPLATE)
    doc.add_heading(title, level=1)
    _add_content(doc, content)

    if sections:
        for block in sections.split("|"):
            parts = block.split(":", 1)
            doc.add_heading(parts[0].strip(), level=2)
            if len(parts) > 1 and parts[1].strip():
                _add_content(doc, parts[1].strip())

    doc.save(output)
    return str(output.resolve())
