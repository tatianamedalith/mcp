from docx import Document


def create_word(title: str, content: str, filename: str)-> str:
    
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.save(filename)

    return f"Word document '{filename}' created successfully."